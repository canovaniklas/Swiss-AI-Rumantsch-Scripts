"""
This script processes different input sources (CSV with URLs, DOCX files, JSON files, PDFs,
plain TXT files, or XML files) and normalizes their content into a consistent JSON-Lines format, gzipped.
It was developed for Romansh.
Each line in the output file will be a JSON object with a standardized schema, including
metadata like filename, language, idiom, and source information.

Output JSON-Lines Schema:
-------------------------
{
  "text": "...",                 // Extracted text content
  "id": "<uuid4>",               // Unique identifier for the record
  "filename": "...",             // Original filename or inferred name (e.g., from URL)
  "language": "<de|roh|unknown>", // Detected or manually specified language (e.g., "de", "roh", "en")
  "language_script": "Latn",     // Script of the language (defaults to "Latn")
  "source": "...",               // Indicates the source type (e.g., "csv_url_fetch", "docx_conversion", "xml_conversion")
  "file_path": "...",            // Original local file path (for local files)
  "dump": "...",                 // Original JSON content as a string (for CSV/URL, JSON inputs)
  "language_score": 0.0,         // Placeholder for language detection score (defaults to 0.0)
  "pii_count": 0,                // Placeholder for Personal Identifiable Information count (defaults to 0)
  "url": "...",                  // Original URL (for CSV/URL inputs)
  "idiom": "..."                 // Detected or manually specified idiom
}

Usage:

-  Processing CSV files (fetching JSON from URLs listed in a CSV):

    python universal_jsonl_builder.py \\
        --type csv \\
        --src "/path/to/your/urls.csv" \\
        --url-column "your_url_column_name" \\
        --out "output_articles.jsonl.gz" \\
        --workers 20 \\
        --language "de" \\
        --idiom "Sursilvan"

    * `--url-column`: The name of the column in your CSV that contains the URLs.
    * `--workers`: (Optional) Number of concurrent downloads (default: 10).
    * `--language`: (Optional) Manually set the language for all records. Overrides auto-detection.
    * `--idiom`: (Optional) Manually set the idiom for all records.

-  Processing DOCX files:

    python universal_jsonl_builder.py \\
        --type docx \\
        --src "/path/to/your/docx_folder" \\
        --out "output_documents.jsonl.gz" \\
        --language "roh" \\
        --idiom "Rumantsch Grischun"

    * `--language`: (Optional) Manually set the language for all records. Overrides filename-based detection.
    * `--idiom`: (Optional) Manually set the idiom for all records. Language is inferred from filename (e.g., `_de.docx`, `_rm.docx`).

-   Processing JSON (or JSON-Lines) files (typically for lexicon-style data):
    
    python universal_jsonl_builder.py \\
        --type json \\
        --src "/path/to/your/json_folder" \\
        --out "output_lexica.jsonl.gz" \\
        --language "roh" \\
        --idiom "Sutsilvan"

    or for specific files:
    
    python universal_jsonl_builder.py \\
        --type json \\
        --src "/path/to/file1.json" "/path/to/file2.json" \\
        --out "output_specific_lexica.jsonl.gz" \\
        --language "roh"

    * `--language`: (Optional) Manually set the language for all records. Defaults to "roh" if not provided, based on the original script's context.
    * `--idiom`: (Optional) Manually set the idiom for all records. Will also try to infer from filename (e.g., files containing "grischun", "sutsilvan").

-  Processing TXT files:
    
    python universal_jsonl_builder.py \\
        --type txt \\
        --src "/path/to/your/txt_folder" \\
        --out "output_texts.jsonl.gz" \\
        --language "de" \\
        --idiom "unknown"

    * `--src`: Path to the directory containing your `.txt` files.
    * `--language`: (Optional) Manually set the language for all records. Overrides filename extension-based detection.
    * `--idiom`: (Optional) Manually set the idiom for all records. Language is inferred from filename extension (e.g., `.de`, `.rm`).

- Processing PDFs (note only processes pdfs which have text content)
    
    python universal_jsonl_builder.py \\
        --type pdf \\
        --src "/path/to/your/pdf_folder" \\
        --out "output_pdfs.jsonl.gz" \\
        --language "de" \\
        --idiom "unknown"

    * `--src`: Path to the directory containing your `.pdf` files.
    * `--language`: (Optional) Manually set the language for all records. Overrides auto-detection.
    * `--idiom`: (Optional) Manually set the idiom for all records.

- Processing XML files:

    python universal_jsonl_builder.py \\
        --type xml \\
        --src "/path/to/your/xml_folder" \\
        --out "output_xml.jsonl.gz" \\
        --language "en" \\
        --idiom "unknown"

    * `--src`: Path to the directory containing your `.xml` files.
    * `--language`: (Optional) Manually set the language for all records. Overrides filename extension-based detection.
    * `--idiom`: (Optional) Manually set the idiom for all records.

Dependencies:
- For CSV/URL processing: `requests`, `langid`, `tqdm`
- For DOCX processing: `docx2txt`, `python-docx`
- For PDF processing: `pdfminer.six`
"""
import argparse
import csv
import gzip
import json
import os
import pathlib
import re
import sys
import uuid
import xml.etree.ElementTree as ET 
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Final, Any, Union

_LANGID_AVAILABLE = False
try:
    import langid  # language detection 
    _LANGID_AVAILABLE = True
except ImportError:
    print("Warning: 'langid' not found. Automatic language detection for CSV/URLs will be disabled.", file=sys.stderr)

_REQUESTS_AVAILABLE = False
try:
    import requests 
    _REQUESTS_AVAILABLE = True
except ImportError:
    print("Warning: 'requests' not found. CSV/URL processing will be disabled.", file=sys.stderr)

_DOCX2TXT_AVAILABLE = False
try:
    import docx2txt 
    _DOCX2TXT_AVAILABLE = True
except ImportError:
    print("Warning: 'docx2txt' not found. Falling back to 'python-docx' for .docx files if available.", file=sys.stderr)

_PYTHON_DOCX_AVAILABLE = False
try:
    from docx import Document 
    _PYTHON_DOCX_AVAILABLE = True
except ImportError:
    print("Warning: 'python-docx' not found. .docx file processing will be disabled if 'docx2txt' is also missing.", file=sys.stderr)

_PDFMINER_AVAILABLE = False
try:
    from pdfminer.high_level import extract_text as extract_pdf_text # PDF text extraction
    _PDFMINER_AVAILABLE = True
except ImportError:
    print("Warning: 'pdfminer.six' not found. PDF file processing will be disabled.", file=sys.stderr)


# default values
DEFAULT_LANGUAGE = "unknown"
DEFAULT_LANGUAGE_SCRIPT = "Latn"
DEFAULT_PII_COUNT = 0
DEFAULT_LANGUAGE_SCORE = 0.0

# language regex detection if it is in the name
_DOCX_SUFFIX_PATTERNS: Final[list[tuple[re.Pattern, str]]] = [
    (re.compile(r"_deu\.docx$", re.IGNORECASE), "de"),
    (re.compile(r"_de\.docx$",  re.IGNORECASE), "de"),
    (re.compile(r"_rom\.docx$", re.IGNORECASE), "roh"),
    (re.compile(r"_rm\.docx$",  re.IGNORECASE), "roh"),
]

# lexica helpers
ENTRY_KEYS: tuple[str, ...] = ("entry", "lemma", "word", "text")

IDIOM_MAP = {
    "grischun":  "Rumantsch Grischun",
    "sutsilvan": "Sutsilvan",
    "sutsilv": "Sutsilvan", # Added for abbreviation
    "surmiran": "Surmiran",
    "sursilvan": "Sursilvan",
    "sursilv": "Sursilvan", # Added for abbreviation
    "puter": "Puter",
    "vallader": "Vallader",
}


def create_base_record(
    text: str,
    filename: str = "",
    file_path: str = "",
    language: str = DEFAULT_LANGUAGE,
    language_script: str = DEFAULT_LANGUAGE_SCRIPT,
    source: str = "",
    url: str = "",
    idiom: str = "",
    dump: str = "",
    language_score: float = DEFAULT_LANGUAGE_SCORE,
    pii_count: int = DEFAULT_PII_COUNT,
    id: str = "", # Added id parameter
) -> dict:
    """Creates a base record with all specified fields, setting defaults for missing ones."""
    return {
        "text": text,
        "id": id if id else str(uuid.uuid4()), # Use provided id or generate new UUID
        "filename": filename,
        "language": language,
        "language_script": language_script,
        "source": source,
        "file_path": file_path,
        "dump": dump,
        "language_score": language_score,
        "pii_count": pii_count,
        "url": url,
        "idiom": idiom,
    }

# CSV/URL Processing 

def detect_language_csv(text: str) -> str:
    """Return 'de', 'roh', or 'unknown' from a chunk of text using langid."""
    if not _LANGID_AVAILABLE:
        return DEFAULT_LANGUAGE
    if not text.strip():
        return DEFAULT_LANGUAGE
    try:
        code, _score = langid.classify(text)
        if code in {"de", "roh"}:
            return code
    except Exception:
        pass
    return DEFAULT_LANGUAGE

def basename_from_url(url: str) -> str:
    """
    rtr.ch URLs look like
        https://www.rtr.ch/article/399199924/json
    We convert that to 399199924.json
    """
    parts = url.rstrip("/").split("/")
    if len(parts) >= 2 and parts[-1] == "json":
        return f"{parts[-2]}.json"
    return os.path.basename(parts[-1] or parts[-2])  # fallback

def fetch_json(url: str, timeout: float = 15.0) -> Union[dict, None]:
    """GET url and return parsed JSON (or None on any failure)."""
    if not _REQUESTS_AVAILABLE:
        return None
    try:
        session = requests.Session()
        session.headers.update({"User-Agent": "bulk-json-fetcher/1.0 (+github.com/...)"})
        r = session.get(url, timeout=timeout)
        r.raise_for_status()
        return r.json()
    except Exception:
        return None

def to_record_csv(url: str, manual_language: str, manual_idiom: str) -> Union[dict, None]:
    """Download, normalise and wrap in the desired schema."""
    data = fetch_json(url)
    if data is None:
        return None

    text = data.get("text") or data.get("body") or json.dumps(data, ensure_ascii=False)
    
    language = manual_language if manual_language else detect_language_csv(text)
    idiom = manual_idiom

    return create_base_record(
        text=text,
        filename=basename_from_url(url),
        language=language,
        language_script=DEFAULT_LANGUAGE_SCRIPT,
        source="csv_url_fetch",
        url=url,
        idiom=idiom,
        dump=json.dumps(data, ensure_ascii=False)
    )

def process_csv_urls(
    csv_path: str,
    url_column: str,
    out_path: str,
    workers: int,
    manual_language: str,
    manual_idiom: str,
):
    if not _REQUESTS_AVAILABLE:
        print("Error: 'requests' library is not installed. Cannot process CSV/URLs.", file=sys.stderr)
        return
    if not _LANGID_AVAILABLE:
        print("Error: 'langid' library is not installed. Cannot perform language detection for CSV/URLs.", file=sys.stderr)
        pass

    try:
        from tqdm import tqdm
    except ImportError:
        print("Warning: 'tqdm' not found. Progress bar will not be displayed.", file=sys.stderr)
        tqdm = lambda x, **kwargs: x 

    with open(csv_path, newline="", encoding="utf-8") as fh:
        reader = csv.DictReader(fh)
        urls = [row[url_column].strip() for row in reader if row.get(url_column)]

    with gzip.open(out_path, "wt", encoding="utf-8") as gz, \
         ThreadPoolExecutor(max_workers=workers) as pool:

        futures = {
            pool.submit(to_record_csv, url, manual_language, manual_idiom): url
            for url in urls
        }
        for f in tqdm(as_completed(futures), total=len(futures), unit="url"):
            rec = f.result()
            if rec is not None:
                gz.write(json.dumps(rec, ensure_ascii=False) + "\n")

    print(f"✅  Wrote {out_path} ({os.path.getsize(out_path):,} bytes)")


# DOCX Processing 

def extract_with_docx2txt(path: pathlib.Path) -> Union[str, None]:
    """Preferred extractor – fast, handles headers/footers/text boxes."""
    if not _DOCX2TXT_AVAILABLE:
        return None
    try:
        return docx2txt.process(str(path))
    except Exception:
        return None


def extract_with_python_docx(path: pathlib.Path) -> str:
    """Fallback extractor using python-docx; includes tables and headers/footers."""
    if not _PYTHON_DOCX_AVAILABLE:
        raise ImportError("Neither 'docx2txt' nor 'python-docx' is installed, cannot process DOCX.")

    doc = Document(path)
    chunks: list[str] = []


    chunks.extend(p.text for p in doc.paragraphs if p.text.strip())

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text.strip():
                    chunks.append(cell.text)

   
    for sect in doc.sections:
        for hf in (sect.header, sect.footer):
            chunks.extend(p.text for p in hf.paragraphs if p.text.strip())

    return "\n".join(chunks)


def docx_to_text(path: pathlib.Path) -> str:
    """Extract raw text – first try docx2txt, then python-docx."""
    text = extract_with_docx2txt(path)
    if text is None:
        text = extract_with_python_docx(path)
    return text


def lang_from_filename_docx(name: str) -> str:
    for pat, lang in _DOCX_SUFFIX_PATTERNS:
        if pat.search(name):
            return lang
    return DEFAULT_LANGUAGE

def process_docx_files(
    src_path: str, out_path: str, manual_language: str, manual_idiom: str
):
    if not _DOCX2TXT_AVAILABLE and not _PYTHON_DOCX_AVAILABLE:
        print("Error: Neither 'docx2txt' nor 'python-docx' is installed. Cannot process DOCX files.", file=sys.stderr)
        return

    src_root = pathlib.Path(src_path).expanduser().resolve()
    if not src_root.is_dir():
        sys.exit(f"Source directory not found: {src_root}")

    files = sorted(src_root.glob("*.docx"))
    if not files:
        sys.exit("No .docx files found in the source directory.")

    out_file_path = pathlib.Path(out_path).expanduser().resolve()
    out_file_path.parent.mkdir(parents=True, exist_ok=True)

    total, skipped = 0, 0
    lang_counts: dict[str, int] = {}

    with gzip.open(out_file_path, "wt", encoding="utf-8") as gz:
        for path in files:
            try:
                text = docx_to_text(path)
                if not text:
                    raise ValueError("Could not extract text from DOCX.")

                language = manual_language if manual_language else lang_from_filename_docx(path.name)
                
                idiom = manual_idiom
                if not idiom and language == 'roh':
                   
                    for idiom_key, idiom_label in IDIOM_MAP.items():
                        if idiom_key in path.name.lower():
                            idiom = idiom_label
                            break

                rec = create_base_record(
                    text=text,
                    filename=path.name,
                    file_path=str(path),
                    language=language,
                    idiom=idiom,
                    source="docx_conversion",
                )
            except Exception as e:
                print(f"[skip] {path.name}: {e}", file=sys.stderr)
                skipped += 1
                continue

            json.dump(rec, gz, ensure_ascii=False)
            gz.write("\n")

            total += 1
            lang_counts[rec["language"]] = (
                lang_counts.get(rec["language"], 0) + 1
            )

    summary = ", ".join(f"{k}:{v}" for k, v in sorted(lang_counts.items()))
    print(
        f"Wrote {out_file_path}  ({total} documents, {skipped} skipped → {summary})",
        file=sys.stderr,
    )

# JSON (or JSONL) Processing 

def detect_idiom_json(fname: str) -> str:
    """Return lower-case idiom key: grischun / sursilvan / sutsilvan."""
    fname_lower = fname.lower()
    for key, label in IDIOM_MAP.items():
        if key in fname_lower:
            return label
    return ""


def extract_terms(obj: Any) -> list[str]:
    """Return all string terms contained anywhere in *obj*."""
    if obj is None:
        return []

    if isinstance(obj, str):
        s = obj.strip()
        return [s] if s else []

    if isinstance(obj, (list, tuple)):
        out: list[str] = []
        for item in obj:
            out.extend(extract_terms(item))
        return out

    if isinstance(obj, dict):
        for key in ENTRY_KEYS:
            if key in obj:
                return extract_terms(obj[key])
        return extract_terms(list(obj.values()))

    return []

def process_json_files(
    src_paths_raw: list[str], out_path: str, manual_language: str, manual_idiom: str
):
    files: list[pathlib.Path] = []
    for p_str in src_paths_raw:
        p = pathlib.Path(p_str).expanduser()
        if p.is_dir():
            files.extend(sorted(p.glob("*.json")))
            files.extend(sorted(p.glob("*.jsonl")))
        elif p.is_file() and p.suffix.lower() in {".json", ".jsonl"}:
            files.append(p)
    
    if not files:
        sys.exit("No JSON or JSON-Lines files found. Nothing to do.")

    out_file_path = pathlib.Path(out_path).expanduser().resolve()
    out_file_path.parent.mkdir(parents=True, exist_ok=True)

    with gzip.open(out_file_path, "wt", encoding="utf-8") as gz:
        for path in files:
            idiom_from_filename = detect_idiom_json(path.name)
            idiom = manual_idiom if manual_idiom else idiom_from_filename

            try:
               
                if path.suffix.lower() == ".jsonl":
                    with path.open(encoding="utf-8") as f:
                        for line in f:
                            data = json.loads(line)
                            process_json_record(data, path, manual_language, idiom, gz)
                else: # .json file
                    data = json.loads(path.read_text(encoding="utf-8"))
                    process_json_record(data, path, manual_language, idiom, gz)

            except json.JSONDecodeError as err:
                print(f"Cannot parse {path}: {err}", file=sys.stderr)
                continue
            except Exception as e:
                print(f"Error processing {path}: {e}", file=sys.stderr)
                continue

    print(f"Wrote {out_file_path}")

def process_json_record(data: Any, path: pathlib.Path, manual_language: str, idiom: str, gz_file_handle: Any):
    """Helper to process a single JSON record/object and write to gz file."""
    terms = extract_terms(data)
    if not terms:
        return

    text_block = "\n".join(terms)
    if idiom:
        text_block = f"This is the {idiom} lexicon:\n" + text_block
    
    language = manual_language if manual_language else "roh"

    record = create_base_record(
        text=text_block,
        filename=path.name,
        file_path=str(path),
        language=language,
        language_script=DEFAULT_LANGUAGE_SCRIPT,
        source="json_lexicon",
        idiom=idiom,
        dump=json.dumps(data, ensure_ascii=False)
    )
    gz_file_handle.write(json.dumps(record, ensure_ascii=False) + "\n")

# TXT Processing 

def detect_language_txt(path: pathlib.Path) -> str:
    """Infer language from filename extension (.de / .rm)."""
    ext = path.suffix.lstrip(".").lower()
    return ext if ext in {"de", "roh"} else DEFAULT_LANGUAGE

def iter_text_files(root: pathlib.Path):
    """Yield all regular files under *root* in lexicographic order."""
    for p in sorted(root.rglob("*")):         
        if p.is_file():
            yield p

def process_txt_files(
    src_path: str, out_path: str, manual_language: str, manual_idiom: str
):
    src_root = pathlib.Path(src_path).expanduser().resolve()
    if not src_root.is_dir():
        sys.exit(f"Source directory not found: {src_root}")

    out_file_path = pathlib.Path(out_path).expanduser().resolve()
    out_file_path.parent.mkdir(parents=True, exist_ok=True)

    with gzip.open(out_file_path, "wt", encoding="utf-8") as fout:
        for path in iter_text_files(src_root):
            try:
                text = path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                try:
                    text = path.read_text(encoding="latin-1")
                except Exception as e:
                    print(f"Warning: Could not read {path.name} with UTF-8 or latin-1. Skipping. Error: {e}", file=sys.stderr)
                    continue
            
            language = manual_language if manual_language else detect_language_txt(path)
            idiom = manual_idiom
            if not idiom and language == 'roh': 
                for idiom_key, idiom_label in IDIOM_MAP.items():
                    if idiom_key in path.name.lower():
                        idiom = idiom_label
                        break


            rec = create_base_record(
                text=text,
                filename=path.name,
                file_path=str(path),
                language=language,
                language_script=DEFAULT_LANGUAGE_SCRIPT,
                source="txt_conversion",
                idiom=idiom
            )
            json.dump(rec, fout, ensure_ascii=False)
            fout.write("\n")

    print(f"Wrote {out_file_path}")

# PDF Processing (NEW)

def detect_language_pdf(path: pathlib.Path) -> str:
    """Infer language from filename extension (.de / .rm) or default."""
    ext = path.suffix.lstrip(".").lower()
    return ext if ext in {"de", "roh"} else DEFAULT_LANGUAGE

def process_pdf_files(
    src_path: str, out_path: str, manual_language: str, manual_idiom: str
):
    if not _PDFMINER_AVAILABLE:
        print("Error: 'pdfminer.six' library is not installed. Cannot process PDF files.", file=sys.stderr)
        return

    src_root = pathlib.Path(src_path).expanduser().resolve()
    if not src_root.is_dir():
        sys.exit(f"Source directory not found: {src_root}")

    files = sorted(src_root.glob("*.pdf"))
    if not files:
        sys.exit("No .pdf files found in the source directory.")

    out_file_path = pathlib.Path(out_path).expanduser().resolve()
    out_file_path.parent.mkdir(parents=True, exist_ok=True)

    total, skipped = 0, 0
    lang_counts: dict[str, int] = {}

    with gzip.open(out_file_path, "wt", encoding="utf-8") as gz:
        for path in files:
            try:
                
                text = extract_pdf_text(str(path))
                if not text.strip(): 
                    raise ValueError("Could not extract any meaningful text from PDF.")

                language = manual_language if manual_language else detect_language_pdf(path)
                idiom = manual_idiom
                if not idiom and language == 'roh':
                    for idiom_key, idiom_label in IDIOM_MAP.items():
                        if idiom_key in path.name.lower():
                            idiom = idiom_label
                            break

                rec = create_base_record(
                    text=text,
                    filename=path.name,
                    file_path=str(path),
                    language=language,
                    idiom=idiom,
                    source="pdf_conversion",
                )
            except Exception as e:
                print(f"[skip] {path.name}: {e}", file=sys.stderr)
                skipped += 1
                continue

            json.dump(rec, gz, ensure_ascii=False)
            gz.write("\n")

            total += 1
            lang_counts[rec["language"]] = (
                lang_counts.get(rec["language"], 0) + 1
            )

    summary = ", ".join(f"{k}:{v}" for k, v in sorted(lang_counts.items()))
    print(
        f"Wrote {out_file_path}  ({total} documents, {skipped} skipped → {summary})",
        file=sys.stderr,
    )

# XML Processing 

def extract_text_from_xml(element: ET.Element) -> str:
    """Recursively extracts all text content from an XML element and its children."""
    texts = []
    if element.text:
        texts.append(element.text.strip())
    for child in element:
        texts.append(extract_text_from_xml(child))
        if child.tail:
            texts.append(child.tail.strip())
    return " ".join(filter(None, texts)).strip()

def process_xml_files(
    src_path: str, out_path: str, manual_language: str, manual_idiom: str
):
    src_root = pathlib.Path(src_path).expanduser().resolve()
    if not src_root.is_dir():
        sys.exit(f"Source directory not found: {src_root}")

    files = sorted(src_root.glob("*.xml"))
    if not files:
        sys.exit("No .xml files found in the source directory.")

    out_file_path = pathlib.Path(out_path).expanduser().resolve()
    out_file_path.parent.mkdir(parents=True, exist_ok=True)

    total, skipped = 0, 0
    lang_counts: dict[str, int] = {}

    with gzip.open(out_file_path, "wt", encoding="utf-8") as gz:
        for path in files:
            try:
                tree = ET.parse(path)
                root = tree.getroot()

                for doc_elem in root.findall(".//DOC"):
                    doc_id = doc_elem.get("id", "")
                    lang_attr = (doc_elem.get("{http://www.w3.org/XML/1998/namespace}lang") or "").strip()

                    idiom_label = manual_idiom or ""
                    if lang_attr:
                        m = re.match(r"(?i)^rm[-_](\w+)", lang_attr)
                        if m:
                            key = m.group(1).lower()
                            idiom_label = IDIOM_MAP.get(key, key.capitalize())

                    text_elem = doc_elem.find("TEXT")
                    doc_text = extract_text_from_xml(text_elem) if text_elem is not None else ""

                    if not doc_text.strip():
                        print(f"[skip] {path.name} (DOC id={doc_id}): No meaningful text extracted. Skipping this DOC.", file=sys.stderr)
                        skipped += 1
                        continue

                    language = manual_language if manual_language else "roh"

                    rec = create_base_record(
                        text=doc_text,
                        filename=path.name,
                        file_path=str(path),
                        language=language,
                        idiom=idiom_label,
                        source="xml_conversion",
                        id=doc_id if doc_id else str(uuid.uuid4()),
                    )
                    json.dump(rec, gz, ensure_ascii=False)
                    gz.write("\n")
                    total += 1
                    lang_counts[language] = lang_counts.get(language, 0) + 1

            except ET.ParseError as e:
                print(f"[skip] {path.name}: XML parsing error ({e}). Skipping file.", file=sys.stderr)
                skipped += 1
                continue
            except Exception as e:
                print(f"[skip] {path.name}: Unexpected error ({e}). Skipping file.", file=sys.stderr)
                skipped += 1
                continue

    summary = ", ".join(f"{k}:{v}" for k, v in sorted(lang_counts.items()))
    print(f"Wrote {out_file_path}  ({total} documents, {skipped} skipped → {summary})", file=sys.stderr)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Convert various data types to a gzipped JSON-Lines corpus."
    )
    parser.add_argument(
        "--type",
        choices=["csv", "docx", "json", "txt", "pdf", "xml"],
        required=True,
        help="Type of input data (csv, docx, json, txt, pdf, xml). 'json' also handles .jsonl files.",
    )
    parser.add_argument(
        "--src",
        required=True,
        nargs="+", 
        help="Source: path to CSV file, directory for DOCX/TXT/PDF/XML, or space-separated paths/directories for JSON/JSONL.",
    )
    parser.add_argument(
        "--out",
        required=True,
        help="Output file name (should end in .jsonl.gz).",
    )
    parser.add_argument(
        "--url-column",
        default="concat",
        help="Column name that holds the URLs (only for --type csv).",
    )
    parser.add_argument(
        "--workers",
        type=int,
        default=10,
        help="Concurrent downloads (only for --type csv, default: 10).",
    )
    parser.add_argument(
        "--language",
        default="",
        help="Manually specify language (e.g., 'de', 'roh', 'en'). Overrides auto-detection.",
    )
    parser.add_argument(
        "--idiom",
        default="",
        help="Manually specify idiom (e.g., 'Rumantsch Grischun', 'Sursilvan').",
    )

    args = parser.parse_args()

    output_dir = pathlib.Path(args.out).parent
    output_dir.mkdir(parents=True, exist_ok=True)

    src_single_path = args.src[0] if len(args.src) == 1 else None

    if args.type == "csv":
        if src_single_path is None:
            sys.exit("For --type csv, --src must be a single path to a CSV file.")
        process_csv_urls(
            src_single_path, args.url_column, args.out, args.workers, args.language, args.idiom
        )
    elif args.type == "docx":
        if src_single_path is None:
            sys.exit("For --type docx, --src must be a single path to a directory.")
        process_docx_files(src_single_path, args.out, args.language, args.idiom)
    elif args.type == "json":
        process_json_files(args.src, args.out, args.language, args.idiom)
    elif args.type == "txt":
        if src_single_path is None:
            sys.exit("For --type txt, --src must be a single path to a directory.")
        process_txt_files(src_single_path, args.out, args.language, args.idiom)
    elif args.type == "pdf":
        if src_single_path is None:
            sys.exit("For --type pdf, --src must be a single path to a directory.")
        process_pdf_files(src_single_path, args.out, args.language, args.idiom)
    elif args.type == "xml": 
        if src_single_path is None:
            sys.exit("For --type xml, --src must be a single path to a directory.")
        process_xml_files(src_single_path, args.out, args.language, args.idiom)